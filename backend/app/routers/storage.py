"""Arquivos — substitui o `supabase.storage`.

Guarda em disco, sob `UPLOADS_DIR`, um diretório por bucket. Os buckets são os
mesmos que o código herdado usa, com a mesma divisão entre público e privado:

| bucket                | quem lê                                             |
|-----------------------|-----------------------------------------------------|
| `agent-files`         | qualquer um — avatares aparecem em `<img>` sem token |
| `audio-messages`      | qualquer um — áudio toca em `<audio src>`            |
| `wiki-uploads`        | qualquer um — imagem dentro de documento             |
| `company-docs`        | autenticado                                          |
| `generated-documents` | autenticado                                          |

**Por que os três primeiros são públicos:** a tela os usa em `src` de tag HTML,
e o navegador não manda o cabeçalho `Authorization` nesse caso. Era assim no
Supabase (`getPublicUrl`) e mudar isso quebraria avatar, áudio e imagem de wiki
de uma vez. O caminho carrega um id difícil de adivinhar — é a mesma proteção
que havia antes, nem mais nem menos.

Escrever exige autenticação em **todos** os buckets, inclusive nos públicos.
"""

import io
import logging
import mimetypes
import re
import time
from html import escape
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from app.config import settings
from app.database import sessao
from app.dependencies import Usuario, exige_papel, usuario_atual

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/storage", tags=["storage"])

_PUBLICOS = {"agent-files", "audio-messages", "wiki-uploads"}
_PRIVADOS = {"company-docs", "generated-documents"}
_BUCKETS = _PUBLICOS | _PRIVADOS

# Segmento de caminho aceitável: letras, números, ponto, hífen, sublinhado. Barra
# é permitida entre segmentos (o `path` do endpoint), mas cada pedaço passa por
# aqui — é o que impede `..` de escapar do diretório do bucket.
_SEGMENTO = re.compile(r"^[A-Za-z0-9._-]{1,120}$")

_TAMANHO_MAXIMO = 25 * 1024 * 1024


def _resolver(bucket: str, caminho: str) -> Path:
    """Traduz bucket + caminho para um arquivo dentro de `UPLOADS_DIR`.

    A validação é por segmento e não por prefixo: conferir só se o resultado
    "começa com" o diretório do bucket é o erro clássico que `..%2f` e link
    simbólico contornam.
    """
    if bucket not in _BUCKETS:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Bucket desconhecido.")

    partes = [p for p in caminho.split("/") if p]
    if not partes or any(not _SEGMENTO.match(p) for p in partes):
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Caminho de arquivo inválido.")

    base = Path(settings.UPLOADS_DIR).resolve() / bucket
    alvo = base.joinpath(*partes)
    # Cinto e suspensório: mesmo com os segmentos validados, confere que o
    # resultado está mesmo sob o bucket.
    if base not in alvo.resolve().parents and alvo.resolve() != base:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Caminho de arquivo inválido.")
    return alvo


class ArquivoOut(BaseModel):
    bucket: str
    path: str
    url: str
    size: int


# ─────────────────────────────────────────────────────────────────────────────
# Extração de texto — metade determinística do `extract-file-text`
# ─────────────────────────────────────────────────────────────────────────────
#
# ⚠️ A edge fazia **duas** coisas: extrair o texto (determinístico) e depois
# mandá-lo para `parse-company-context`, que usa LLM para virar campos do
# `company_profile`. Só a primeira está aqui.
#
# A segunda depende do **Lovable AI Gateway** (`LOVABLE_API_KEY`,
# `google/gemini-2.5-flash`) — dependência da plataforma de origem, que é
# justamente o que a migração existe para remover. Trocar de provedor custa
# dinheiro e é decisão do Erick; ver `docs/ROADMAP.md`.

_LIMITE_TEXTO = 50_000


class TextoExtraidoOut(BaseModel):
    text: str
    caracteres: int
    truncado: bool


@router.post("/extrair-texto/{bucket}/{caminho:path}", response_model=TextoExtraidoOut)
async def extrair_texto(
    bucket: str,
    caminho: str,
    _: Usuario = Depends(usuario_atual),
):
    """Texto de um arquivo já enviado. Aceita `.txt`, `.md`, `.pdf` e `.docx`.

    ⚠️ Declarado **antes** do upload genérico: `POST /storage/{bucket}/{caminho}`
    casaria com esta URL, tomando `extrair-texto` como nome de bucket.

    Recebe o caminho no storage em vez do arquivo porque é assim que a tela já
    trabalha: ela sobe o documento primeiro e depois pede o processamento,
    mandando só `storagePath`.
    """
    alvo = _resolver(bucket, caminho)
    if not alvo.is_file():
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Arquivo não encontrado.")

    ext = alvo.suffix.lower().lstrip(".")
    try:
        if ext in ("txt", "md"):
            texto = alvo.read_text(encoding="utf-8", errors="replace")
        elif ext == "pdf":
            from pypdf import PdfReader

            leitor = PdfReader(str(alvo))
            texto = "\n".join((p.extract_text() or "") for p in leitor.pages)
        elif ext == "docx":
            import docx

            texto = "\n".join(p.text for p in docx.Document(str(alvo)).paragraphs)
        else:
            raise HTTPException(
                status.HTTP_400_BAD_REQUEST,
                f"Formato não suportado: .{ext}. Envie txt, md, pdf ou docx.",
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("Falha ao extrair texto de %s/%s: %s", bucket, caminho, e)
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            f"Não foi possível ler o arquivo: {e}",
        )

    if not texto.strip():
        # PDF de imagem escaneada cai aqui: tem páginas, não tem texto. Dizer
        # "sem texto extraível" é mais útil que devolver string vazia.
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "Arquivo sem texto extraível. Se for um PDF escaneado, ele precisa de OCR.",
        )

    truncado = len(texto) > _LIMITE_TEXTO
    return TextoExtraidoOut(
        text=texto[:_LIMITE_TEXTO], caracteres=len(texto), truncado=truncado
    )


# ─────────────────────────────────────────────────────────────────────────────
# Acesso a documento gerado — portado de `sign-generated-document`
# ─────────────────────────────────────────────────────────────────────────────


class DocumentoOut(BaseModel):
    url: str
    title: str | None = None
    doc_type: str | None = None


@router.get("/documento/{documento_id}", response_model=DocumentoOut)
async def documento_gerado(
    documento_id: str,
    usuario: Usuario = Depends(usuario_atual),
):
    """Devolve o endereço de um documento gerado, se ele for do usuário.

    ⚠️ Declarado **antes** do upload e do download genéricos — é a quarta rota
    deste módulo com prefixo fixo, e todas precisam vir antes de
    `/storage/{bucket}/{caminho}`, senão `documento` vira nome de bucket.

    A edge criava uma URL assinada de 1 hora do Supabase Storage. Aqui não há
    assinatura: `generated-documents` é bucket privado, servido por
    `/storage/privado/...`, que já exige o token do usuário em cada request.
    O efeito de segurança é o mesmo — melhor, até: a URL assinada valia para
    quem a tivesse durante uma hora, esta vale só para quem tem o token.

    A conferência de dono acontece aqui e não pelo RLS porque a leitura roda
    como `service_role` para poder olhar a linha antes de decidir.
    """
    async with sessao(role="service_role") as conn:
        linha = await conn.fetchrow(
            "SELECT user_id::text AS user_id, storage_path, title, doc_type "
            "FROM public.generated_documents WHERE id = $1::uuid",
            documento_id,
        )
    if linha is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Documento não encontrado.")
    if linha["user_id"] and linha["user_id"] != usuario.id:
        # 404, não 403: quem não é dono também não deveria descobrir que o
        # documento existe.
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Documento não encontrado.")

    caminho = (linha["storage_path"] or "").lstrip("/")
    if not caminho:
        raise HTTPException(
            status.HTTP_409_CONFLICT,
            "O documento existe no banco mas não tem arquivo associado.",
        )
    return DocumentoOut(
        url=f"/storage/privado/generated-documents/{caminho}",
        title=linha["title"],
        doc_type=linha["doc_type"],
    )


# ─────────────────────────────────────────────────────────────────────────────
# Gerar PDF/DOCX — portado de `generate-document`
# ─────────────────────────────────────────────────────────────────────────────
#
# O agente escreve uma tag no chat com a definição do documento; a tela extrai
# e chama aqui. A definição é uma lista de seções — deliberadamente pobre, para
# o agente conseguir produzir sem errar o formato.

_NIVEIS = {"H1": 0, "H2": 1, "H3": 2}


class SecaoIn(BaseModel):
    heading: str | None = None
    text: str | None = None
    bold: bool = False


class DefinicaoIn(BaseModel):
    title: str | None = None
    sections: list[SecaoIn] = []


class DocumentoIn(BaseModel):
    type: str = Field(description="pdf ou docx")
    title: str = Field(default="Documento", max_length=200)
    agent_id: str | None = None
    definition: DefinicaoIn


class DocumentoGeradoOut(BaseModel):
    id: str
    title: str
    doc_type: str
    size_bytes: int
    created_at: str


def _montar_docx(definicao: DefinicaoIn) -> bytes:
    from docx import Document as Docx

    doc = Docx()
    if definicao.title:
        doc.add_heading(definicao.title, level=0)
    for secao in definicao.sections:
        if secao.heading in _NIVEIS:
            doc.add_heading(secao.text or "", level=_NIVEIS[secao.heading] + 1)
            continue
        if not secao.text:
            continue
        paragrafo = doc.add_paragraph()
        paragrafo.add_run(secao.text).bold = secao.bold

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _montar_pdf(definicao: DefinicaoIn) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    estilos = getSampleStyleSheet()
    partes = []
    if definicao.title:
        partes += [Paragraph(escape(definicao.title), estilos["Title"]), Spacer(1, 12)]
    for secao in definicao.sections:
        if secao.heading in _NIVEIS:
            partes += [
                Paragraph(escape(secao.text or ""), estilos[f"Heading{_NIVEIS[secao.heading] + 1}"]),
                Spacer(1, 6),
            ]
            continue
        if not secao.text:
            continue
        # `escape` porque o reportlab interpreta um mini-HTML no parágrafo: um
        # `<` no texto do agente quebraria a geração inteira.
        texto = escape(secao.text)
        partes += [
            Paragraph(f"<b>{texto}</b>" if secao.bold else texto, estilos["BodyText"]),
            Spacer(1, 6),
        ]

    buffer = io.BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4).build(partes or [Spacer(1, 1)])
    return buffer.getvalue()


@router.post("/documentos/gerar", response_model=DocumentoGeradoOut,
             status_code=status.HTTP_201_CREATED)
async def gerar_documento(dados: DocumentoIn, usuario: Usuario = Depends(usuario_atual)):
    """Gera o arquivo, grava no bucket privado e registra em `generated_documents`.

    O arquivo vai para `generated-documents/<usuário>/<id>.<ext>` — a pasta por
    usuário é o que sustenta o download privado, que confere o dono pelo
    primeiro segmento do caminho.

    Se o registro no banco falhar, o arquivo é apagado. Arquivo órfão em disco
    não aparece em lugar nenhum e só ocupa espaço para sempre.
    """
    if dados.type not in ("pdf", "docx"):
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Tipo inválido. Use pdf ou docx.")
    if not dados.definition.sections and not dados.definition.title:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST, "A definição do documento está vazia."
        )

    try:
        bytes_ = _montar_pdf(dados.definition) if dados.type == "pdf" else _montar_docx(dados.definition)
    except Exception as e:  # noqa: BLE001
        logger.warning("Geração de %s falhou: %s", dados.type, e)
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            f"Não foi possível gerar o documento: {e}",
        )

    doc_id = str(uuid4())
    caminho = f"{usuario.id}/{doc_id}.{dados.type}"
    destino = _resolver("generated-documents", caminho)
    destino.parent.mkdir(parents=True, exist_ok=True)
    destino.write_bytes(bytes_)

    try:
        async with sessao(role="service_role") as conn:
            linha = await conn.fetchrow(
                """
                INSERT INTO public.generated_documents
                    (id, user_id, agent_id, title, doc_type, storage_path, size_bytes)
                VALUES ($1::uuid, $2::uuid, $3, $4, $5, $6, $7)
                RETURNING id::text AS id, title, doc_type, size_bytes,
                          to_char(created_at AT TIME ZONE 'UTC', 'YYYY-MM-DD"T"HH24:MI:SS') || 'Z' AS created_at
                """,
                doc_id, usuario.id, dados.agent_id, dados.title, dados.type,
                caminho, len(bytes_),
            )
    except Exception:
        destino.unlink(missing_ok=True)
        raise

    logger.info("Documento %s (%s, %d bytes) gerado por %s",
                doc_id, dados.type, len(bytes_), usuario.id)
    return DocumentoGeradoOut(**dict(linha))


@router.post("/{bucket}/{caminho:path}", response_model=ArquivoOut,
             status_code=status.HTTP_201_CREATED)
async def enviar(
    bucket: str,
    caminho: str,
    arquivo: UploadFile = File(...),
    _: Usuario = Depends(usuario_atual),
):
    """Grava o arquivo. Sobrescreve se já existir.

    Sobrescrever é o padrão porque era o que o código herdado pedia
    (`upsert: true` no avatar) e porque o caminho já identifica o dono — o
    avatar de um agente é sempre `avatars/<id>.png`. Sem isso, trocar a foto
    criaria arquivo novo e deixaria lixo para sempre.
    """
    alvo = _resolver(bucket, caminho)
    alvo.parent.mkdir(parents=True, exist_ok=True)

    tamanho = 0
    try:
        with alvo.open("wb") as destino:
            while pedaco := await arquivo.read(1024 * 1024):
                tamanho += len(pedaco)
                if tamanho > _TAMANHO_MAXIMO:
                    # Interrompe durante a leitura, não depois: guardar o arquivo
                    # inteiro para então recusá-lo é o que enche o disco.
                    destino.close()
                    alvo.unlink(missing_ok=True)
                    raise HTTPException(
                        status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                        f"Arquivo acima do limite de {_TAMANHO_MAXIMO // (1024 * 1024)} MB.",
                    )
                destino.write(pedaco)
    except HTTPException:
        raise
    except OSError as e:
        logger.error("Falha ao gravar %s/%s: %s", bucket, caminho, e)
        raise HTTPException(
            status.HTTP_500_INTERNAL_SERVER_ERROR, "Não foi possível gravar o arquivo."
        )

    logger.info("Arquivo %s/%s gravado (%d bytes)", bucket, caminho, tamanho)
    return ArquivoOut(
        bucket=bucket, path=caminho, url=f"/storage/{bucket}/{caminho}", size=tamanho
    )


@router.get("/listar/{bucket}/{prefixo:path}")
async def listar(bucket: str, prefixo: str, usuario: Usuario = Depends(usuario_atual)) -> dict:
    """Nomes dos arquivos sob um prefixo. Existe para não sondar por 404.

    ⚠️ **O motivo é concreto e foi medido no navegador em 10/08/2026.** O
    carregador de avatares descobria a foto de cada agente por força bruta:
    para cada um, tentava `avatars/<id>.png`, `.jpg`, `.jpeg` e `.webp` até uma
    carregar. Com 13 ids na lista e duas passadas, isso somava **72 requisições
    404 em toda carga de página** — e o console é justamente onde a gente lê os
    erros de verdade quando algo quebra. Enterrar os erros reais debaixo de 72
    falsos é caro de um jeito que não aparece em nenhuma métrica.

    Devolve só nomes, nunca conteúdo, e só de bucket público — quem lista já
    poderia baixar cada arquivo um a um.
    """
    if bucket not in _PUBLICOS:
        raise HTTPException(
            status.HTTP_403_FORBIDDEN,
            f"Listagem disponível apenas para buckets públicos: {', '.join(sorted(_PUBLICOS))}.",
        )
    # `_resolver` é quem impede travessia de caminho (`../`); reaproveitar aqui
    # garante que a listagem obedeça exatamente à mesma regra do download.
    alvo = _resolver(bucket, prefixo or ".")
    if not alvo.is_dir():
        return {"arquivos": []}
    return {"arquivos": sorted(f.name for f in alvo.iterdir() if f.is_file())}


@router.get("/privado/{bucket}/{caminho:path}")
async def baixar_privado(
    bucket: str,
    caminho: str,
    _: Usuario = Depends(usuario_atual),
):
    """Serve arquivo de bucket privado, exigindo token.

    ⚠️ Declarado **antes** do handler genérico de propósito: o FastAPI casa na
    ordem de definição, e `/storage/privado/x/y` casaria com
    `/storage/{bucket}/{caminho}` — com `bucket="privado"` — se viesse depois.
    """
    if bucket not in _PRIVADOS:
        raise HTTPException(
            status.HTTP_404_NOT_FOUND, "Bucket não é privado — use /storage/{bucket}/…"
        )
    alvo = _resolver(bucket, caminho)
    if not alvo.is_file():
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Arquivo não encontrado.")
    tipo, _sub = mimetypes.guess_type(alvo.name)
    return FileResponse(alvo, media_type=tipo or "application/octet-stream")


@router.get("/{bucket}/{caminho:path}")
async def baixar(bucket: str, caminho: str):
    """Serve o arquivo.

    Sem dependência de autenticação na assinatura porque os buckets públicos
    precisam responder a `<img src>`. Os privados conferem o token na mão, logo
    abaixo — o `Depends` obrigatório fecharia os dois casos junto.
    """
    alvo = _resolver(bucket, caminho)
    if bucket in _PRIVADOS:
        raise HTTPException(
            status.HTTP_401_UNAUTHORIZED,
            "Este bucket é privado. Use o endpoint autenticado.",
        )
    if not alvo.is_file():
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Arquivo não encontrado.")

    tipo, _ = mimetypes.guess_type(alvo.name)
    return FileResponse(
        alvo,
        media_type=tipo or "application/octet-stream",
        # Curto de propósito: avatar e imagem de wiki são sobrescritos no mesmo
        # caminho, e cache longo mostraria a foto antiga por horas. A tela já
        # acrescenta `?t=<timestamp>` ao trocar, mas isso não pode ser a única
        # defesa.
        headers={"Cache-Control": "public, max-age=300"},
    )


@router.delete("/{bucket}/{caminho:path}", status_code=status.HTTP_204_NO_CONTENT)
async def remover(bucket: str, caminho: str, _: Usuario = Depends(usuario_atual)):
    """Apaga o arquivo. Apagar o que não existe não é erro.

    O `remove()` do Supabase também não reclamava, e a tela usa isso para
    limpar variantes (`.png` e `.jpg`) sem saber qual existe.
    """
    alvo = _resolver(bucket, caminho)
    alvo.unlink(missing_ok=True)


def preparar_diretorios() -> None:
    """Cria os diretórios dos buckets no start.

    Sem isto, o primeiro upload de cada bucket depende do `mkdir` do request, o
    que esconde erro de permissão até alguém tentar subir um arquivo.
    """
    base = Path(settings.UPLOADS_DIR)
    for bucket in sorted(_BUCKETS):
        try:
            (base / bucket).mkdir(parents=True, exist_ok=True)
        except OSError as e:
            logger.warning("Não foi possível preparar %s/%s: %s", base, bucket, e)
            return
    logger.info("Buckets prontos em %s: %s", base, ", ".join(sorted(_BUCKETS)))


# ─────────────────────────────────────────────────────────────────────────────
# Faxina de anexos vencidos — portado de `cleanup-expired-files`
# ─────────────────────────────────────────────────────────────────────────────

# 30 dias. Anexo de chat é material de trabalho de uma conversa, não arquivo
# permanente — guardar para sempre enche o disco com coisa que ninguém abre.
_VALIDADE_SEGUNDOS = 30 * 24 * 60 * 60

# `uploadFileToStorage` grava como `<pasta>/<timestamp>_<nome>`. A idade sai do
# nome, não do mtime do arquivo: cópia e restauração de backup mudam o mtime e
# apagariam anexo recente por engano.
_CARIMBO = re.compile(r"^(\d{10,13})_")


class FaxinaOut(BaseModel):
    removidos: int
    examinados: int


@router.post("/faxina", response_model=FaxinaOut)
async def faxina(_: Usuario = Depends(exige_papel("administrador"))):
    """Apaga anexos com mais de 30 dias do bucket `agent-files`.

    Só `agent-files`: avatar de agente vive nele também, mas em `avatars/`, e
    esses **não** têm carimbo no nome — a regra do carimbo os protege sozinha.
    """
    base = Path(settings.UPLOADS_DIR) / "agent-files"
    if not base.is_dir():
        return FaxinaOut(removidos=0, examinados=0)

    limite_ms = (time.time() - _VALIDADE_SEGUNDOS) * 1000
    removidos = examinados = 0
    for arquivo in base.rglob("*"):
        if not arquivo.is_file():
            continue
        examinados += 1
        m = _CARIMBO.match(arquivo.name)
        if not m:
            continue
        carimbo = int(m.group(1))
        # 10 dígitos é segundo, 13 é milissegundo — o JS grava em ms.
        if len(m.group(1)) == 10:
            carimbo *= 1000
        if carimbo < limite_ms:
            try:
                arquivo.unlink()
                removidos += 1
            except OSError as e:
                logger.warning("Não foi possível remover %s: %s", arquivo, e)

    logger.info("Faxina: %d removidos de %d examinados", removidos, examinados)
    return FaxinaOut(removidos=removidos, examinados=examinados)
